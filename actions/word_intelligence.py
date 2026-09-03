"""Word & Document Intelligence for JARVIS.

Provides structural understanding of Microsoft Word documents:
- Finding and navigating to headings ("Put this under the heading called Results")
- Section and paragraph placement ("Put this underneath the second paragraph", "Continue where I left off")
- Safe paragraph/section replacement with destructive action confirmations
- Document structure inspection and outlines without dumping entire documents into AI context
- Powered by Windows COM UI Automation with python-docx fallback.
"""

from __future__ import annotations

import os
import platform
import sys
import time
from pathlib import Path
from typing import Any, Callable, Optional

_SYSTEM = platform.system()


class WordDocumentIntelligence:
    """Intelligent document navigator and structural editor for Microsoft Word."""

    @staticmethod
    def is_windows() -> bool:
        return _SYSTEM == "Windows"

    @classmethod
    def get_word_app(cls) -> Optional[Any]:
        """Obtain active Word.Application COM instance on Windows."""
        if not cls.is_windows():
            return None
        try:
            import win32com.client
            return win32com.client.GetActiveObject("Word.Application")
        except Exception:
            return None

    @classmethod
    def get_active_document(cls) -> Optional[Any]:
        app = cls.get_word_app()
        if app:
            try:
                return app.ActiveDocument
            except Exception:
                return None
        return None

    @classmethod
    def get_document_outline(cls, doc_path: Optional[str] = None) -> dict:
        """Inspect and return the outline/headings of the active document without bloated tokens."""
        headings = []
        para_count = 0
        doc_title = "Untitled"

        if cls.is_windows():
            doc = cls.get_active_document()
            if doc:
                try:
                    doc_title = doc.Name
                    para_count = doc.Paragraphs.Count
                    # Scan paragraphs for headings (Word style 'Heading 1', 'Heading 2', etc.)
                    for i in range(1, min(para_count + 1, 100)):
                        try:
                            para = doc.Paragraphs(i)
                            text = para.Range.Text.strip()
                            style_name = str(para.Style.NameLocal)
                            if "Heading" in style_name or "Başlık" in style_name or (len(text) < 80 and para.Range.Bold):
                                headings.append({
                                    "index": i,
                                    "text": text,
                                    "style": style_name,
                                })
                        except Exception:
                            continue
                    return {
                        "ok": True,
                        "title": doc_title,
                        "total_paragraphs": para_count,
                        "headings": headings,
                    }
                except Exception as e:
                    print(f"[WordIntelligence] COM Outline error: {e}")

        # Fallback to python-docx if a document path was specified or COM unavailable
        if doc_path and Path(doc_path).exists():
            try:
                import docx
                doc = docx.Document(doc_path)
                for idx, p in enumerate(doc.paragraphs, start=1):
                    if p.style and "Heading" in p.style.name:
                        headings.append({"index": idx, "text": p.text.strip(), "style": p.style.name})
                return {
                    "ok": True,
                    "title": Path(doc_path).name,
                    "total_paragraphs": len(doc.paragraphs),
                    "headings": headings,
                }
            except Exception as e:
                return {"ok": False, "error": f"Failed to parse document: {e}"}

        return {
            "ok": True,
            "title": "Simulated Document",
            "total_paragraphs": 5,
            "headings": [{"index": 1, "text": "Introduction", "style": "Heading 1"},
                         {"index": 4, "text": "Results", "style": "Heading 1"}],
        }

    @classmethod
    def navigate_and_insert(
        cls,
        target_heading: Optional[str] = None,
        paragraph_index: Optional[int] = None,
        text_to_type: str = "",
        speed: str = "human",
        speak: Optional[Callable[[str], None]] = None,
    ) -> str:
        """Locate target heading or paragraph, position cursor, and type the text with human cadence."""
        from actions.word_typing import WindowsWordController, type_text_humanlike

        if not cls.is_windows():
            # In simulated environment
            typed = type_text_humanlike(text_to_type, speed="instant")
            loc = f"heading '{target_heading}'" if target_heading else f"paragraph {paragraph_index}"
            return f"Successfully typed {typed} characters under {loc} (simulated environment)."

        app = cls.get_word_app()
        if not app:
            return "Microsoft Word isn't open and I couldn't locate the document."

        doc = cls.get_active_document()
        if not doc:
            return "No active document open in Microsoft Word."

        # Bring Word window to foreground safely
        word_windows = WindowsWordController.get_word_windows()
        if word_windows:
            WindowsWordController.focus_word_window(word_windows[0]["hwnd"])

        placed = False
        target_desc = ""

        # Locate by heading text
        if target_heading:
            target_desc = f"heading '{target_heading}'"
            try:
                find_obj = doc.Content.Find
                find_obj.ClearFormatting()
                find_obj.Text = target_heading
                find_obj.Forward = True
                find_obj.Wrap = 1  # wdFindContinue

                if find_obj.Execute():
                    found_range = find_obj.Parent
                    # Position selection right after the heading and start a new paragraph
                    app.Selection.SetRange(found_range.End, found_range.End)
                    app.Selection.TypeParagraph()
                    placed = True
            except Exception as e:
                print(f"[WordIntelligence] Failed to find heading via COM: {e}")

        # Locate by paragraph index
        elif paragraph_index is not None and paragraph_index > 0:
            target_desc = f"paragraph {paragraph_index}"
            try:
                count = doc.Paragraphs.Count
                idx = min(max(1, paragraph_index), count)
                target_para = doc.Paragraphs(idx)
                # Position selection at the end of the target paragraph
                app.Selection.SetRange(target_para.Range.End, target_para.Range.End)
                placed = True
            except Exception as e:
                print(f"[WordIntelligence] Failed to position at paragraph: {e}")

        if not placed:
            # Default to end of document if specific placement could not be matched
            target_desc = "end of document (target location not found)"
            try:
                app.Selection.EndKey(Unit=6)  # wdStory
            except Exception:
                pass

        if speak:
            speak(f"Positioned insertion point at {target_desc}. Now typing, sir.")

        # Type using authentic keystroke automation
        typed = type_text_humanlike(
            text_to_type,
            speed=speed,
            safety_check=WindowsWordController.is_word_foreground,
        )

        return f"Successfully placed and typed {typed} characters under {target_desc} in Word."

    @classmethod
    def replace_section(
        cls,
        target_text_or_heading: str,
        new_text: str,
        confirmed: bool = False,
        speed: str = "human",
        speak: Optional[Callable[[str], None]] = None,
    ) -> str:
        """Safely replace text or section in Microsoft Word.

        Enforces Section 12 rule: requires confirmation for destructive/replacement operations.
        """
        if not target_text_or_heading:
            return "No target section or text specified to replace."

        # Safety confirmation enforcement
        if not confirmed:
            return (
                f"SAFETY CONFIRMATION REQUIRED: Are you sure you want to replace '{target_text_or_heading}' "
                f"with the new text in Microsoft Word? Please confirm to proceed with this replacement."
            )

        if not cls.is_windows():
            return f"Confirmed. Replaced '{target_text_or_heading}' with new text (simulated environment)."

        app = cls.get_word_app()
        if not app or not cls.get_active_document():
            return "Microsoft Word isn't open and I couldn't locate the document."

        doc = cls.get_active_document()
        try:
            find_obj = doc.Content.Find
            find_obj.ClearFormatting()
            find_obj.Text = target_text_or_heading
            find_obj.Forward = True

            if find_obj.Execute():
                found_range = find_obj.Parent
                # Select the target range and replace it
                found_range.Select()
                # Type new content
                from actions.word_typing import type_text_humanlike, WindowsWordController
                typed = type_text_humanlike(
                    new_text,
                    speed=speed,
                    safety_check=WindowsWordController.is_word_foreground,
                )
                return f"Successfully replaced '{target_text_or_heading}' with {typed} typed characters."
            else:
                return f"Could not find '{target_text_or_heading}' in the active Word document."
        except Exception as e:
            return f"Error replacing section in Word: {e}"


def word_intelligence(
    parameters: dict,
    player=None,
    speak: Optional[Callable[[str], None]] = None,
) -> str:
    """Main tool entry point for Word Document Intelligence."""
    action = parameters.get("action", "insert").lower().strip()
    heading = parameters.get("heading") or parameters.get("target_heading")
    paragraph = parameters.get("paragraph") or parameters.get("paragraph_index")
    text = parameters.get("text") or parameters.get("content") or ""
    speed = parameters.get("speed", "human")
    confirmed = bool(parameters.get("confirmed", False) or parameters.get("force", False))

    if paragraph is not None:
        try:
            paragraph = int(paragraph)
        except ValueError:
            paragraph = None

    if action in ("outline", "structure", "headings"):
        res = WordDocumentIntelligence.get_document_outline()
        headings_str = ", ".join([h["text"] for h in res.get("headings", [])]) or "None found"
        return f"Document '{res.get('title')}': {res.get('total_paragraphs')} paragraphs. Headings: {headings_str}."

    elif action in ("replace", "rewrite_section"):
        target = heading or parameters.get("target_text") or ""
        return WordDocumentIntelligence.replace_section(
            target_text_or_heading=target,
            new_text=text,
            confirmed=confirmed,
            speed=speed,
            speak=speak,
        )

    else:  # "insert", "continue", "write_under"
        if not text:
            return "No text provided to insert into Word."
        return WordDocumentIntelligence.navigate_and_insert(
            target_heading=heading,
            paragraph_index=paragraph,
            text_to_type=text,
            speed=speed,
            speak=speak,
        )
